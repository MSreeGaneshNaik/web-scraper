"""
Extract readable text from downloadable files (PDF, Word, Excel).

Each function is best-effort: if the library is missing or the file is
malformed, it returns an empty string rather than raising.
"""

import io


def extract_pdf(data):
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""
    try:
        reader = PdfReader(io.BytesIO(data))
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return "\n".join(parts).strip()
    except Exception:
        return ""


def extract_docx(data):
    try:
        import docx
    except ImportError:
        return ""
    try:
        doc = docx.Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()
    except Exception:
        return ""


def extract_xlsx(data):
    try:
        import openpyxl
    except ImportError:
        return ""
    try:
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for ws in wb.worksheets:
            parts.append(f"# Sheet: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()
    except Exception:
        return ""


def extract_document(url, data, content_type=""):
    """Pick an extractor based on the URL extension / content type."""
    u = url.lower()
    ct = (content_type or "").lower()
    if u.endswith(".pdf") or "pdf" in ct:
        return extract_pdf(data)
    if u.endswith(".docx") or "wordprocessingml" in ct:
        return extract_docx(data)
    if u.endswith(".xlsx") or "spreadsheetml" in ct:
        return extract_xlsx(data)
    return ""
